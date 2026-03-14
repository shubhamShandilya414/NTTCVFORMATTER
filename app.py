"""
CV Transformer App
==================
Converts any CV/Resume into the NTT template format using:
- HuggingFace sentence-transformers for field matching
- Groq LLaMA-3.3-70b-versatile for content extraction
- python-docx for output generation
"""

import streamlit as st
import tempfile
import os
import json
import re
import numpy as np
from pathlib import Path

# ─── Page config ───
st.set_page_config(
    page_title="CV Transformer — NTT DATA",
    page_icon="🔷",
    layout="wide",
)

# ─── Custom CSS ───
st.markdown("""
<style>
    /* NTT DATA Brand Colors */
    :root {
        --ntt-blue: #0072BC;
        --ntt-dark: #002E6D;
        --ntt-light: #E8F1FA;
        --ntt-logo-blue: #6585C2;
        --ntt-accent: #00A5E5;
    }
    .main-header {
        background: linear-gradient(135deg, #002E6D 0%, #0072BC 60%, #00A5E5 100%);
        padding: 2rem 2rem 1.5rem 2rem;
        border-radius: 12px;
        margin-bottom: 1.5rem;
        text-align: center;
        border-bottom: 4px solid #00A5E5;
    }
    .main-header h1 { color: #FFFFFF; margin: 0; font-size: 2.2rem; letter-spacing: 0.5px; }
    .main-header p  { color: #B8D4F0; margin: 0.4rem 0 0 0; font-size: 1rem; }
    .main-header .dev-credit { color: #7BBEEE; font-size: 0.78rem; margin-top: 0.6rem; opacity: 0.9; }
    .step-card {
        background: #E8F1FA;
        border-left: 4px solid #0072BC;
        padding: 1rem 1.2rem;
        border-radius: 0 8px 8px 0;
        margin-bottom: 1rem;
    }
    .step-card h3 { margin: 0 0 0.3rem 0; color: #002E6D; }
    .match-good  { color: #27ae60; font-weight: 600; }
    .match-ok    { color: #f39c12; font-weight: 600; }
    .match-poor  { color: #e74c3c; font-weight: 600; }
    .stDownloadButton > button {
        background: linear-gradient(135deg, #0072BC, #002E6D) !important;
        color: white !important;
        border: none !important;
        padding: 0.6rem 2rem !important;
        font-size: 1.1rem !important;
    }
    .stDownloadButton > button:hover {
        background: linear-gradient(135deg, #00A5E5, #0072BC) !important;
    }
    /* Primary button styling */
    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, #0072BC, #002E6D) !important;
        border: none !important;
    }
    /* Info boxes */
    .stAlert { border-left-color: #0072BC !important; }
    /* Sidebar */
    section[data-testid="stSidebar"] {
        background: #f0f5fb;
        border-right: 2px solid #0072BC;
    }
    .footer-box {
        text-align: center;
        padding: 1rem;
        background: linear-gradient(135deg, #002E6D, #0072BC);
        border-radius: 8px;
        margin-top: 1rem;
    }
    .footer-box p { color: #B8D4F0; font-size: 0.8rem; margin: 0; }
    .footer-box .dev { color: #FFFFFF; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.3rem; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
    <h1>📄 CV Transformer — NTT DATA Format</h1>
    <p>Upload any CV → AI matches fields & extracts content → Download in NTT template format</p>
    <p class="dev-credit">Developed by Shubham Shandilya</p>
</div>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════
#  UTILITY FUNCTIONS
# ═══════════════════════════════════════════════════════════════

@st.cache_resource(show_spinner="Loading HuggingFace embedding model …")
def load_embedding_model():
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")


def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from a PDF using pdfplumber (or fallback to PyPDF2)."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    except ImportError:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_path)
        return "\n".join(p.extract_text() or "" for p in reader.pages)


def extract_text_from_docx(docx_path: str) -> str:
    from docx import Document
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def extract_template_sections(docx_path: str) -> list:
    """Parse the NTT template and return its section structure."""
    from docx import Document
    doc = Document(docx_path)
    sections = []
    current = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if _is_template_heading(para, text):
            if current:
                sections.append(current)
            current = {"heading": text, "content": ""}
        else:
            if current is None:
                current = {"heading": "__NAME_HEADER__", "content": ""}
            current["content"] += text + "\n"

    if current:
        sections.append(current)

    # Also note table structures
    for ti, table in enumerate(doc.tables):
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        sections.append({
            "heading": f"__TABLE_{ti}__",
            "content": " | ".join(header_cells),
            "is_table": True,
            "headers": header_cells,
        })

    return sections


def _is_template_heading(para, text):
    style = para.style.name.lower() if para.style else ""
    if "heading" in style:
        return True
    if len(text.split()) <= 5 and para.runs:
        all_bold = all(r.bold for r in para.runs if r.text.strip())
        if all_bold and not text.startswith("Client:") and not text.startswith("Role:"):
            return True
    return False


def parse_input_cv_sections(full_text: str) -> list:
    """
    Heuristic section splitter for raw CV text.
    Looks for lines that are likely section headings.
    """
    lines = full_text.split("\n")
    sections = []
    current = None
    heading_keywords = {
        # Summary / Objective
        "objective", "summary", "profile", "professional summary",
        "profile summary", "career objective", "job objective",
        "career summary", "executive summary", "about me",
        # Experience
        "experience", "work experience", "professional experience",
        "employment", "work history", "detailed experience",
        "previous work experience", "key result areas",
        "employment history", "career history",
        # Education
        "education", "qualifications", "academic qualifications",
        "academic background",
        # Skills
        "skills", "technical skills", "it skills", "key skills",
        "core skills", "tools and technologies",
        # Competencies (CRITICAL: these were being missed)
        "competencies", "core competencies", "technical competencies",
        "key competencies", "functional competencies",
        # Soft Skills (CRITICAL: was being missed)
        "soft skills", "interpersonal skills", "behavioral skills",
        # Achievements (CRITICAL: was being missed)
        "achievements", "key achievements", "accomplishments",
        "awards", "honors", "recognition", "awards and achievements",
        # Certifications
        "certifications", "training", "training & certification",
        "training and certification", "licenses", "credentials",
        "professional certifications",
        # Projects
        "projects", "key projects",
        # Specializations
        "specializations", "specializations & implementations",
        # Others
        "publications", "references", "languages", "hobbies",
        "interests", "contact", "contact details", "personal",
        "personal details", "additional information",
    }

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        is_heading = False
        lower = stripped.lower().rstrip(":")
        # Direct match against known headings
        if lower in heading_keywords:
            is_heading = True
        # ALL CAPS short text (e.g., "WORK EXPERIENCE", "CORE COMPETENCIES")
        elif stripped.isupper() and 2 < len(stripped) < 60 and len(stripped.split()) <= 7:
            is_heading = True
        # Title Case multi-word known heading (e.g., "Professional Summary:")
        elif lower in heading_keywords:
            is_heading = True
        # Fuzzy: check if any known keyword is a substring (e.g., "Key Achievements & Awards")
        elif not is_heading and len(stripped.split()) <= 7:
            for kw in heading_keywords:
                if kw in lower and len(lower) < 50:
                    is_heading = True
                    break

        if is_heading:
            if current:
                sections.append(current)
            current = {"heading": stripped.rstrip(":"), "content": ""}
        else:
            if current is None:
                current = {"heading": "__HEADER__", "content": ""}
            current["content"] += stripped + "\n"

    if current:
        sections.append(current)
    return sections


def compute_field_mapping(input_sections, template_sections, model, threshold=0.25):
    """Use HuggingFace embeddings + cosine similarity to match fields."""
    from sklearn.metrics.pairwise import cosine_similarity

    # Enriched representations
    input_texts = []
    for s in input_sections:
        preview = s["content"][:300] if s["content"] else ""
        input_texts.append(f'{s["heading"]}: {preview}')

    template_texts = []
    # NTT template sections + sub-categories that feed into Skills
    ntt_fields = [
        "Name and Country header personal information contact details email phone",
        "Professional Summary career objective profile overview about me executive summary",
        "Key Achievements accomplishments awards recognition metrics impact results numbers percentages revenue",
        "Education academic qualifications degrees university college school",
        "Training and Certifications certifications professional certifications licenses credentials courses AWS certified scrum master azure",
        "Skills IT Skills technical skill sets programming languages frameworks tools databases OS cloud products DevOps IDE build tools technical competencies",
        "Core Competencies key competencies functional competencies domain expertise business competencies areas of expertise",
        "Soft Skills interpersonal skills behavioral skills communication leadership teamwork negotiation problem solving",
        "Professional Experience organization and duration employment history company timeline",
        "Detailed Experience with client role project responsibilities work details job duties specializations implementations",
    ]
    ntt_labels = [
        "Personal Info / Name",
        "Professional Summary",
        "Key Achievements",
        "Education",
        "Training & Certifications",
        "Skills / IT Skills",
        "Skills → Core Competencies",
        "Skills → Soft Skills",
        "Professional Experience",
        "Detailed Experience",
    ]
    template_headings = []
    for s in template_sections:
        h = s["heading"]
        if h.startswith("__TABLE_") or h.startswith("__NAME_"):
            continue
        template_texts.append(f'{h}: {s["content"][:200]}')
        template_headings.append(h)

    # Also add the canonical NTT fields for better matching
    all_template_texts = template_texts + ntt_fields
    all_template_labels = template_headings + ntt_labels

    if not input_texts or not all_template_texts:
        return {}

    input_emb = model.encode(input_texts, convert_to_numpy=True, show_progress_bar=False)
    templ_emb = model.encode(all_template_texts, convert_to_numpy=True, show_progress_bar=False)

    sim_matrix = cosine_similarity(input_emb, templ_emb)

    # For each input section, find best template match
    mapping = []
    for i, sec in enumerate(input_sections):
        best_j = int(np.argmax(sim_matrix[i]))
        best_score = float(sim_matrix[i][best_j])
        target_label = all_template_labels[best_j] if best_j < len(all_template_labels) else "Unknown"
        mapping.append({
            "input_heading": sec["heading"],
            "input_content": sec["content"],
            "matched_template_field": target_label,
            "similarity": round(best_score, 4),
        })

    return mapping, sim_matrix


def call_groq_extraction(api_key: str, full_cv_text: str, field_mapping: list) -> dict:
    """
    Use Groq LLaMA-3.3-70b-versatile to extract and restructure CV content
    into the NTT template JSON structure.
    """
    from groq import Groq
    client = Groq(api_key=api_key)

    mapping_summary = "\n".join(
        f'  - Input "{m["input_heading"]}" → Template "{m["matched_template_field"]}" (sim={m["similarity"]})'
        for m in field_mapping
    )

    prompt = f"""You are an expert CV/Resume data extractor. Given raw CV text and a field mapping,
extract ALL information EXACTLY as written and restructure it into the NTT template JSON format.

The NTT template has EXACTLY these sections and NOTHING else. Only return these fields:

## TARGET NTT TEMPLATE STRUCTURE:

1. **name**: Full name of the candidate (string)
2. **country**: Country (e.g., "INDIA") (string)

3. **professional_summary**: A CONCISE professional summary paragraph (3-5 sentences, max 80 words).
   
   RULES for professional_summary:
   - If the CV has a dedicated "Professional Summary" or "Profile Summary" section with 1-3 sentences,
     copy it VERBATIM.
   - If the CV opens with a LONG bullet-point list (like 5-10 bullets describing the candidate),
     DO NOT copy them all raw. Instead, synthesize the key highlights into a SHORT flowing paragraph
     of 3-5 sentences. Extract the essence: years of experience, domain, key expertise, current role.
   - NEVER include raw bullet characters (•, ▪, -) in this field.
   - NEVER exceed 5 sentences or 80 words. Be concise.
   - DO NOT include achievements, competencies, soft skills, or certifications here.
   
   Example of a GOOD summary for a 17-year payments expert:
   "Senior Product Consultant with 17+ years of IT experience, including 13+ years in financial services 
   and payments domain. Subject matter expert on ACI Money Transfer System and Fiserv Dovetail with 
   extensive implementation experience across SWIFT, SEPA, FedWire, and real-time payment schemes. 
   Currently working at LTIMindtree on wholesale payment solutions."

4. **key_achievements**: List of achievement strings from "Key Achievements" or "Achievements" section.
   *** Copy each achievement VERBATIM — preserve ALL numbers, metrics, percentages, revenue figures. ***
   *** Example: "Successfully managed 3 large-scale projects simultaneously, overseeing 40 resources & driving revenues exceeding €10 million" must be kept EXACTLY. ***
   If no achievements section exists in the CV, return empty list [].

5. **education**: List of education entries. Each entry:
   {{"degree": "...", "institution": "...", "year": "..."}}
   Include ALL education entries found (degrees, diplomas, executive programs, etc.)

6. **certifications**: List of training and certification strings.
   Extract EXACTLY as written from the CV. Look in sections like "Certifications",
   "Training & Certification", "Training", "Licenses", "Professional Certifications".
   Also look for certifications mentioned ANYWHERE in the CV (e.g., in profile summary or skills).
   Examples: "AWS Certified Cloud Practitioner – September 2022",
             "Certified Scrum Master (CSM)",
             "Senior Management Executive Program (1 year), IIM Kozhikode, 2022",
             "Azure Fundamentals - AZ 900"
   *** Include the full certification name with date/institution if available. ***
   *** If no certifications or training found anywhere in the CV, return EMPTY list []. ***

7. **skills_table**: Technical/IT skills classified into categories.
   Format: [{{"category": "Languages", "value": "Java, Python, SQL"}}, ...]
   
   Available categories and what goes in each:
   - "Languages": ONLY actual programming/scripting languages (Java, Python, C, C++, SQL, COBOL, Groovy, JavaScript, etc.)
   - "Frameworks": Software frameworks (Spring Boot, Hibernate, JUnit, React, Django, etc.)
   - "Database": Database systems (Oracle, PostgreSQL, MySQL, SQL Server, DB2, MongoDB, etc.)
   - "O.S.": Operating systems (Windows, Linux, Unix, AIX, z/OS, etc.)
   - "Cloud": Cloud platforms & services (AWS Lambda, API Gateway, Azure, GCP, etc.)
   - "CI/CD Tools": CI/CD tools (Jenkins, SonarQube, GitHub Actions, etc.)
   - "DevOps Tools": DevOps tools (Jenkins, GitHub, Bitbucket, Docker, Kubernetes, etc.)
   - "Development IDE's": IDEs (Eclipse, IntelliJ, Toad, PL/SQL Developer, etc.)
   - "Build Tools": Build tools (Maven, Gradle, Ant, etc.)
   - "WebServices": Web service types (REST, SOAP, GraphQL, etc.)
   - "Tools": Misc tools (JIRA, Postman, GIT, Confluence, etc.)
   - "Products": Payment/banking/domain-specific PRODUCTS and PLATFORMS
     (ACI Money Transfer System, Fiserv Dovetail, Oracle FLEXCUBE, NetReveal, MuleSoft Anypoint, etc.)

   *** MISCLASSIFICATION DENY-LIST — these are NOT programming languages or frameworks: ***
   *** SWIFT (payment messaging system) → goes in "Products", NOT "Languages" ***
   *** SEPA (Single Euro Payments Area) → goes in "Products", NOT "Frameworks" ***
   *** CHIPS, FedWire, CHAPS, BACS, Zelle, Interac, FedNow, STET → "Products" ***
   *** ISO 20022, ISO 8583, CBPR+, MT/MX → "Products" ***
   *** OFAC, AML, KYC → NOT IT skills at all, skip them ***
   *** Agile, Scrum, Waterfall, SAFe → NOT IT skills, skip them ***
   *** WebSphere MQ / IBM MQ → "Tools", NOT "Languages" ***

   *** RULES: ***
   *** - ONLY include a category if the candidate ACTUALLY has skills listed for it. ***
   *** - If the value would be empty, DO NOT include that category. ***
   *** - If the CV has NO technical/IT skills at all, return EMPTY list []. ***
   *** - Do NOT fabricate skills. Only extract what is EXPLICITLY stated. ***

8. **core_competencies**: List of domain/business competency strings found in the CV.
   Examples: "Payment Product Development", "Business Analysis", "Functional Scope Documentation",
   "Product Management", "RFP and RFI Response handling", "Wire & Real Time Payments"
   Extract EXACTLY as written from sections like "Core Competencies", "Key Competencies",
   "Specializations", "Specializations & Implementations", "Areas of Expertise", etc.
   Also extract from opening bullet-point highlights that describe domain expertise.
   If the CV has no such content, return empty list [].

9. **soft_skills**: List of soft/interpersonal skill strings found in the CV.
   Examples: "Team Building", "Critical Thinking", "Negotiation", "Decision-Making"
   Extract EXACTLY as written from sections like "Soft Skills", "Interpersonal Skills", etc.
   If the CV has no such section, return empty list [].

10. **professional_experience**: List of ALL organizations with durations:
   [{{"organization": "FIS Global", "duration": "Jan 2014 to Present"}}, ...]
   Include ALL organizations, including previous/earlier/brief ones.

11. **detailed_experience**: List of ALL project/role details across the ENTIRE career:
   [{{
     "client": "Company or Client name — MUST NEVER be empty",
     "role": "Job title",
     "duration": "Start - End date",
     "project_details": "Project name or brief description of work done",
     "responsibilities": ["list of bullet points"]
   }}]

   *** CRITICAL RULES for detailed_experience: ***

   A) "client" must NEVER be empty. ALWAYS fill it with the company/organization name.
      - "IEXCEED TECHNOLOGIES as Project Manager" → client = "IEXCEED TECHNOLOGIES"
      - "FERROGREEN TECHNOLOGIES as Project Engineer" → client = "FERROGREEN TECHNOLOGIES"

   B) Brief one-line entries in "Previous Work Experience" MUST still get a full entry:
      "Feb'13-Jan'14: IEXCEED TECHNOLOGIES as Project Manager" →
      {{"client": "IEXCEED TECHNOLOGIES", "role": "Project Manager",
        "duration": "Feb 2013 - Jan 2014", "project_details": "", "responsibilities": []}}

   C) If one company has MULTIPLE sub-roles at different clients, create SEPARATE entries:
      HCL TECHNOLOGIES with sub-roles at Rabobank, Bank of Tokyo Mitsubishi →
      separate entry for each with client = "HCL Technologies - Rabobank", etc.

   D) Copy each responsibility bullet EXACTLY from the CV. NEVER summarize.
      KEEP ALL metrics: numbers, percentages, currencies (€10 million, 40 resources, 10%)

   E) Include EVERY role from the ENTIRE CV — even brief/old entries at the bottom.

## FIELD MAPPING (AI-matched input sections → template sections):
{mapping_summary}

## FULL CV TEXT (extract from the ENTIRE text, including the very end):
---
{full_cv_text}
---

RULES:
1. Return EXACTLY these 11 fields: name, country, professional_summary, key_achievements, education, certifications, skills_table, core_competencies, soft_skills, professional_experience, detailed_experience.
2. professional_summary must be a CONCISE flowing paragraph (3-5 sentences, max 80 words). No bullet points, no bullet characters (•, ▪, -). If the CV opens with a long bullet list, synthesize it into a short paragraph.
3. key_achievements is a SEPARATE list — extract each achievement as a standalone string, preserving ALL metrics and numbers VERBATIM. If no achievements section exists, return empty list [].
4. certifications is a SEPARATE field — do NOT fold it into professional_summary or education.
5. PRESERVE original wording, ALL numbers, metrics, percentages, currencies.
6. Scan the ENTIRE CV to the very last line. Do NOT stop early.
7. Return ONLY valid JSON. No markdown, no code fences, no explanation.
"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": "You are a precise CV data extractor. For professional_summary: write a concise 3-5 sentence paragraph (max 80 words), no bullet points. For all other fields: extract content VERBATIM, never lose metrics/numbers/percentages. Classify skills into proper categories. Return ONLY valid JSON."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.1,
        max_tokens=16000,
    )

    raw = response.choices[0].message.content.strip()
    # Clean up any markdown fences
    raw = re.sub(r'^```json\s*', '', raw)
    raw = re.sub(r'^```\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)

    return json.loads(raw)


def validate_and_fix_skills(data: dict) -> dict:
    """
    Post-processing safety net: fix common LLM misclassifications in skills_table.
    Moves payment/domain terms OUT of Languages/Frameworks and INTO Products.
    Removes non-IT items entirely.
    """
    skills = data.get("skills_table", [])
    if not isinstance(skills, list):
        return data

    # Terms that are PRODUCTS/PLATFORMS, not languages or frameworks
    product_terms = {
        "swift", "sepa", "chips", "fedwire", "fed wire", "chaps", "bacs",
        "zelle", "interac", "fednow", "stet", "iso 20022", "iso20022",
        "iso 8583", "iso8583", "cbpr+", "cbpr", "mt/mx", "mt", "mx",
        "aci money transfer", "aci mts", "aci universal payments",
        "aci upf", "fiserv dovetail", "dovetail", "oracle flexcube",
        "flexcube", "netreveal", "detica", "indus lending",
        "mulesoft", "mulesoft anypoint", "anypoint",
        "websphere mq", "ibm mq", "swift alliance",
        "npp", "ucm", "gpi",
    }

    # Terms that are NOT IT skills at all — remove entirely
    non_it_terms = {
        "agile", "scrum", "waterfall", "safe", "kanban",
        "ofac", "aml", "kyc", "bsa", "fatca",
        "six sigma", "itil", "prince2", "pmp",
        "stakeholder management", "team leadership",
        "project management", "program management",
    }

    # Categories that should NEVER contain payment/domain terms
    protected_categories = {"languages", "frameworks", "webservices"}

    misplaced_products = []  # Collect terms that need to move to Products

    new_skills = []
    for skill in skills:
        if not isinstance(skill, dict):
            continue
        cat = str(skill.get("category", "")).strip()
        val = str(skill.get("value", "")).strip()
        if not cat or not val:
            continue

        cat_lower = cat.lower()

        if cat_lower in protected_categories:
            # Split comma-separated values and check each one
            items = [v.strip() for v in val.split(",")]
            keep_items = []
            for item in items:
                item_lower = item.lower().strip()
                if not item_lower:
                    continue
                # Check if this item is a product term
                is_product = False
                for pt in product_terms:
                    if pt == item_lower or item_lower.startswith(pt) or pt in item_lower:
                        is_product = True
                        misplaced_products.append(item)
                        break
                # Check if it's a non-IT term
                is_non_it = False
                for nt in non_it_terms:
                    if nt == item_lower or item_lower.startswith(nt):
                        is_non_it = True
                        break
                if not is_product and not is_non_it:
                    keep_items.append(item)

            if keep_items:
                skill["value"] = ", ".join(keep_items)
                new_skills.append(skill)
            # If no items left, skip this category entirely
        else:
            # For non-protected categories (Tools, Products, Database, etc.) — keep as-is
            # but still strip out non-IT terms
            items = [v.strip() for v in val.split(",")]
            clean_items = []
            for item in items:
                item_lower = item.lower().strip()
                is_non_it = any(nt == item_lower or item_lower.startswith(nt) for nt in non_it_terms)
                if not is_non_it and item.strip():
                    clean_items.append(item)
            if clean_items:
                skill["value"] = ", ".join(clean_items)
                new_skills.append(skill)

    # If we collected misplaced products, add/merge them into a Products row
    if misplaced_products:
        existing_products = None
        for skill in new_skills:
            if skill.get("category", "").lower() == "products":
                existing_products = skill
                break

        if existing_products:
            existing_vals = [v.strip() for v in existing_products["value"].split(",")]
            for mp in misplaced_products:
                if mp not in existing_vals:
                    existing_vals.append(mp)
            existing_products["value"] = ", ".join(existing_vals)
        else:
            new_skills.append({
                "category": "Products",
                "value": ", ".join(misplaced_products),
            })

    data["skills_table"] = new_skills
    return data


def generate_ntt_docx(data: dict, template_path: str) -> bytes:
    """
    Generate a .docx in NTT format using the docx-js approach via python-docx.
    Replicates the NTT template structure with extracted data.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    import zipfile
    from io import BytesIO

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    # ── Extract NTT logo from the template ──
    logo_bytes = None
    try:
        with zipfile.ZipFile(template_path, 'r') as z:
            for name in z.namelist():
                if name.startswith("word/media/") and ("image" in name.lower()):
                    logo_bytes = z.read(name)
                    break
    except Exception:
        logo_bytes = None

    # ── Helper functions ──
    def add_heading_text(text, size=16, bold=True, color=RGBColor(0, 0, 0), space_before=12, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = 'Calibri'
        return p

    def add_normal_text(text, size=12, bold=False, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.name = 'Calibri'
        return p

    def add_bullet(text, size=11):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        return p

    def add_separator():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # Add a bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '6',
            qn('w:space'): '1',
            qn('w:color'): '000000',
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ═══════════════════════════════════════
    #  0. NTT DATA LOGO
    # ═══════════════════════════════════════
    if logo_bytes:
        try:
            logo_stream = BytesIO(logo_bytes)
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_logo.paragraph_format.space_after = Pt(8)
            run = p_logo.add_run()
            run.add_picture(logo_stream, width=Inches(1.7))
        except Exception:
            pass  # Skip logo if insertion fails

    # ═══════════════════════════════════════
    #  1. NAME & COUNTRY
    # ═══════════════════════════════════════
    name = data.get("name", "Candidate Name")
    country = data.get("country", "INDIA")

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_name.paragraph_format.space_after = Pt(2)
    run = p_name.add_run(name)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Calibri'

    p_country = doc.add_paragraph()
    p_country.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_country.paragraph_format.space_after = Pt(6)
    run = p_country.add_run(country)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'

    add_separator()

    # ═══════════════════════════════════════
    #  2. PROFESSIONAL SUMMARY
    # ═══════════════════════════════════════
    add_heading_text("Professional Summary:", size=12, color=RGBColor(0, 0, 0))
    summary = data.get("professional_summary", "")
    if summary:
        add_normal_text(summary, size=12)

    # ── Key Achievements (only if present in CV) ──
    achievements = data.get("key_achievements", [])
    if isinstance(achievements, list):
        achievements = [str(a).strip() for a in achievements if str(a).strip()]
    else:
        achievements = []

    if achievements:
        add_heading_text("Key Achievements:", size=12, space_before=8)
        for ach in achievements:
            add_bullet(ach, size=11)

    add_separator()

    # ═══════════════════════════════════════
    #  3. EDUCATION
    # ═══════════════════════════════════════
    add_heading_text("Education", size=12)
    education = data.get("education", [])
    if isinstance(education, list):
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree", "")
                inst = edu.get("institution", "")
                year = edu.get("year", "")
                line = f"{degree} from {inst}" + (f" - {year}" if year else "")
                add_bullet(line, size=12)
            elif isinstance(edu, str):
                add_bullet(edu, size=12)
    elif isinstance(education, str):
        add_normal_text(education, size=12)

    add_separator()

    # ═══════════════════════════════════════
    #  3b. TRAINING & CERTIFICATIONS (only if present)
    # ═══════════════════════════════════════
    certifications = data.get("certifications", [])
    if isinstance(certifications, list):
        certifications = [str(c).strip() for c in certifications if str(c).strip()]
    else:
        certifications = []

    if certifications:
        add_heading_text("Training & Certifications", size=12)
        for cert in certifications:
            add_bullet(cert, size=12)
        add_separator()

    # ═══════════════════════════════════════
    #  4. SKILLS (adapts to profile type)
    # ═══════════════════════════════════════
    skills = data.get("skills_table", [])
    core_comp = data.get("core_competencies", [])
    soft_skills = data.get("soft_skills", [])

    # Filter out any skill entries with empty/blank values
    if isinstance(skills, list):
        skills = [
            s for s in skills
            if isinstance(s, dict) and str(s.get("value", "")).strip() and str(s.get("category", "")).strip()
        ]
    else:
        skills = []

    # Clean core_competencies and soft_skills
    if isinstance(core_comp, list):
        core_comp = [str(c).strip() for c in core_comp if str(c).strip()]
    else:
        core_comp = []
    if isinstance(soft_skills, list):
        soft_skills = [str(s).strip() for s in soft_skills if str(s).strip()]
    else:
        soft_skills = []

    has_tech_skills = len(skills) > 0
    has_competencies = len(core_comp) > 0 or len(soft_skills) > 0

    if has_tech_skills or has_competencies:
        add_heading_text("Skills", size=12)

        # Determine table header and sub-heading based on profile type
        if has_tech_skills:
            add_heading_text("IT SKILLS", size=12, space_before=4)
            table_header = "TECHNICAL SKILL SETS"
        else:
            table_header = "SKILL SETS"

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr = table.rows[0]
        for i, txt in enumerate([table_header, table_header]):
            cell = hdr.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            shading = cell._element.get_or_add_tcPr()
            sh_elem = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): 'D9E2F3',
                qn('w:val'): 'clear',
            })
            shading.append(sh_elem)

        def add_skill_row(category, value):
            row = table.add_row()
            row.cells[0].text = str(category)
            row.cells[1].text = str(value)
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(11)
                        r.font.name = 'Calibri'

        # Add technical skills rows
        for skill in skills:
            add_skill_row(skill.get("category", ""), skill.get("value", ""))

        # Add core competencies as a row
        if core_comp:
            add_skill_row("Core Competencies", ", ".join(core_comp))

        # Add soft skills as a row
        if soft_skills:
            add_skill_row("Soft Skills", ", ".join(soft_skills))

        add_separator()

    # ═══════════════════════════════════════
    #  5. PROFESSIONAL EXPERIENCE TABLE
    # ═══════════════════════════════════════
    add_heading_text("Professional Experience", size=12)

    prof_exp = data.get("professional_experience", [])
    if prof_exp and isinstance(prof_exp, list):
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0]
        for i, txt in enumerate(["Name of Organization", "Duration"]):
            cell = hdr.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            shading = cell._element.get_or_add_tcPr()
            sh_elem = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): 'D9E2F3',
                qn('w:val'): 'clear',
            })
            shading.append(sh_elem)

        for exp in prof_exp:
            if isinstance(exp, dict):
                org = exp.get("organization", "")
                dur = exp.get("duration", "")
            elif isinstance(exp, (list, tuple)) and len(exp) >= 2:
                org, dur = exp[0], exp[1]
            else:
                continue
            row = table.add_row()
            row.cells[0].text = str(org)
            row.cells[1].text = str(dur)
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(11)
                        r.font.name = 'Calibri'

    add_separator()

    # ═══════════════════════════════════════
    #  6. DETAILED EXPERIENCE
    # ═══════════════════════════════════════
    add_heading_text("Detailed Experience", size=12)

    detailed = data.get("detailed_experience", [])
    if isinstance(detailed, list):
        for idx, exp in enumerate(detailed):
            if not isinstance(exp, dict):
                continue

            client = exp.get("client", "").strip()
            role = exp.get("role", "").strip()
            duration = exp.get("duration", "").strip()
            project = exp.get("project_details", "").strip()
            responsibilities = exp.get("responsibilities", [])

            # Skip completely empty entries
            if not client and not role:
                continue

            # Client line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"Client: {client}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Calibri'

            # Role + Duration line
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"Role: {role}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Calibri'
            if duration:
                run = p.add_run(f"\t\t\t\t{duration}")
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Calibri'

            # Project Details
            if project:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"Project Details: {project}")
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Calibri'

            # Responsibilities
            if responsibilities:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run("Responsibilities:")
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Calibri'

                if isinstance(responsibilities, list):
                    for resp in responsibilities:
                        add_bullet(str(resp), size=11)
                elif isinstance(responsibilities, str):
                    for line in responsibilities.split("\n"):
                        line = line.strip().lstrip("•-– ")
                        if line:
                            add_bullet(line, size=11)

            if idx < len(detailed) - 1:
                doc.add_paragraph()  # Spacer between entries

    # ── Save to bytes ──
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════
#  STREAMLIT UI
# ═══════════════════════════════════════════════════════════════

# ── Sidebar ──
with st.sidebar:
    st.markdown(
        '<div style="text-align:center; padding:0.5rem 0 1rem 0;">'
        '<span style="font-size:1.4rem; font-weight:700; color:#0072BC; letter-spacing:1px;">NTT DATA</span><br>'
        '<span style="font-size:0.75rem; color:#6585C2;">CV Transformation Tool</span>'
        '</div>',
        unsafe_allow_html=True,
    )
    st.markdown("---")
    st.markdown("### 🧠 How It Works")
    st.markdown("""
    **Step 1** — Enter your Groq API key
    
    **Step 2** — Upload the NTT template (.docx) and an input CV (.docx or .pdf)
    
    **Step 3** — HuggingFace `all-MiniLM-L6-v2` embeddings match CV fields to template
    
    **Step 4** — Groq `LLaMA-3.3-70b-versatile` extracts and restructures content
    
    **Step 5** — A new .docx is generated in NTT DATA format
    """)
    st.markdown("---")
    st.markdown("### ⚙️ Advanced Settings")
    sim_threshold = st.slider("Similarity Threshold", 0.1, 0.9, 0.25, 0.05,
                              help="Minimum cosine similarity for field matching")
    st.markdown("---")
    st.markdown(
        '<div style="text-align:center; color:#6585C2; font-size:0.75rem; padding-top:0.5rem;">'
        'Developed by <strong>Shubham Shandilya</strong>'
        '</div>',
        unsafe_allow_html=True,
    )

# ── Main Area ──

# ── API Key (prominent, top of main area) ──
st.markdown('<div class="step-card"><h3>🔑 Step 1: Groq API Key</h3>Required for LLaMA content extraction</div>',
            unsafe_allow_html=True)
key_col1, key_col2 = st.columns([3, 1])
with key_col1:
    groq_key = st.text_input(
        "Enter your Groq API Key",
        type="password",
        placeholder="gsk_xxxxxxxxxxxxxxxxxxxxxxxx",
        help="Get a free key at https://console.groq.com/keys API Keys → Create",
    )
with key_col2:
    st.markdown("<br>", unsafe_allow_html=True)
    st.link_button("🔗 Get Free Key", "https://console.groq.com/keys", use_container_width=True)

st.markdown("")  # spacer

# ── File uploads ──
col1, col2 = st.columns(2)

with col1:
    st.markdown('<div class="step-card"><h3>📋 Step 2: NTT Template</h3>Upload your target format</div>',
                unsafe_allow_html=True)
    template_file = st.file_uploader("Upload NTT Template (.docx)", type=["docx"], key="template")

with col2:
    st.markdown('<div class="step-card"><h3>📥 Step 3: Input CV</h3>Upload the CV to transform</div>',
                unsafe_allow_html=True)
    input_file = st.file_uploader("Upload Input CV (.docx or .pdf)", type=["docx", "pdf"], key="input_cv")

# ── Processing ──
if template_file and input_file and groq_key:
    st.markdown("---")

    if st.button("🚀 Transform CV", type="primary", use_container_width=True):

        # Save uploads to temp files
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
            tf.write(template_file.getvalue())
            template_path = tf.name

        input_suffix = ".pdf" if input_file.name.lower().endswith(".pdf") else ".docx"
        with tempfile.NamedTemporaryFile(suffix=input_suffix, delete=False) as tf:
            tf.write(input_file.getvalue())
            input_path = tf.name

        try:
            # ─── STEP 1: Parse template ───
            with st.status("🔍 Analyzing template and CV...", expanded=True) as status:
                st.write("Parsing NTT template structure...")
                template_sections = extract_template_sections(template_path)
                st.write(f"  ✅ Found {len(template_sections)} template sections")

                # ─── STEP 2: Extract input CV text ───
                st.write("Extracting text from input CV...")
                if input_suffix == ".pdf":
                    full_text = extract_text_from_pdf(input_path)
                else:
                    full_text = extract_text_from_docx(input_path)
                st.write(f"  ✅ Extracted {len(full_text)} characters")

                input_sections = parse_input_cv_sections(full_text)
                st.write(f"  ✅ Identified {len(input_sections)} sections in input CV")

                # ─── STEP 3: HuggingFace Embedding matching ───
                st.write("Loading HuggingFace embedding model...")
                model = load_embedding_model()

                st.write("Computing cosine similarity for field matching...")
                field_mapping, sim_matrix = compute_field_mapping(
                    input_sections, template_sections, model, threshold=sim_threshold
                )
                st.write(f"  ✅ Matched {sum(1 for m in field_mapping if m['similarity'] > sim_threshold)} fields")
                status.update(label="✅ Analysis complete", state="complete")

            # ─── Display field mapping ───
            st.markdown("### 🔗 Field Mapping (HuggingFace Embeddings)")
            mapping_cols = st.columns([3, 1, 3])
            mapping_cols[0].markdown("**Input CV Section**")
            mapping_cols[1].markdown("**Similarity**")
            mapping_cols[2].markdown("**→ NTT Template Field**")

            for m in field_mapping:
                c1, c2, c3 = st.columns([3, 1, 3])
                c1.write(m["input_heading"])
                score = m["similarity"]
                if score >= 0.5:
                    cls = "match-good"
                elif score >= 0.3:
                    cls = "match-ok"
                else:
                    cls = "match-poor"
                c2.markdown(f'<span class="{cls}">{score:.2f}</span>', unsafe_allow_html=True)
                c3.write(m["matched_template_field"])

            # ─── STEP 4: Groq LLaMA extraction ───
            with st.status("🤖 Extracting with Groq LLaMA-3.3-70b-versatile...", expanded=True) as status:
                st.write("Sending CV to Groq API for intelligent extraction...")
                extracted_data = call_groq_extraction(groq_key, full_text, field_mapping)
                st.write("  ✅ Content extracted and structured")

                st.write("Validating skill classifications...")
                extracted_data = validate_and_fix_skills(extracted_data)
                st.write("  ✅ Skills validated and corrected")

                st.write("Generating NTT format .docx...")
                docx_bytes = generate_ntt_docx(extracted_data, template_path)
                st.write("  ✅ Document generated!")
                status.update(label="✅ Transformation complete!", state="complete")

            # ─── Display extracted JSON preview ───
            with st.expander("📊 Extracted Data Preview (JSON)", expanded=False):
                st.json(extracted_data)

            # ─── Download ───
            st.markdown("---")
            st.markdown("### 📥 Download Transformed CV")

            candidate_name = extracted_data.get("name", "Transformed_CV").replace(" ", "_")
            filename = f"{candidate_name}_NTT_Format.docx"

            st.download_button(
                label=f"⬇️ Download {filename}",
                data=docx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

            st.success("✅ CV successfully transformed to NTT format!")

        except json.JSONDecodeError as e:
            st.error(f"❌ Failed to parse Groq response as JSON. The LLM may have returned invalid format.\n\nError: {e}")
            st.info("Tip: Try again — LLaMA responses can occasionally be malformed.")
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            st.exception(e)
        finally:
            # Cleanup temp files
            for p in [template_path, input_path]:
                try:
                    os.unlink(p)
                except:
                    pass

elif not groq_key:
    st.warning("🔑 Please enter your **Groq API Key** above (Step 1) to get started. [Get a free key →](https://console.groq.com/keys)")
elif not template_file:
    st.info("📋 Please upload the **NTT Template** (.docx) file in Step 2.")
elif not input_file:
    st.info("📥 Please upload an **Input CV** (.docx or .pdf) in Step 3 to transform.")

# ── Footer ──
st.markdown("---")
st.markdown(
    '<div class="footer-box">'
    '<p class="dev">Developed by Shubham Shandilya</p>'
    '<p>Powered by HuggingFace Transformers • Groq LLaMA • Streamlit • python-docx</p>'
    '</div>',
    unsafe_allow_html=True,
)